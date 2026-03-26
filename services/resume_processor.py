"""
services/resume_processor.py — Extração de texto de arquivos de currículo.

Suporta:
  - PDF  → via PyPDF2
  - DOCX → via python-docx

O texto extraído é usado pelo ai_local_service para gerar o match score
e personalizar o currículo com base na vaga.
"""
import os
import uuid
from werkzeug.datastructures import FileStorage
from flask import current_app

# Importação condicional para não quebrar se a lib não estiver instalada
try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def allowed_file(filename: str) -> bool:
    """Verifica se a extensão do arquivo é permitida."""
    allowed = current_app.config.get('ALLOWED_EXTENSIONS', {'pdf', 'docx'})
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extrai texto de um arquivo PDF usando PyPDF2.
    Percorre todas as páginas e concatena o texto.
    """
    if not HAS_PDF:
        raise ImportError("PyPDF2 não está instalado. Execute: pip install PyPDF2")

    text_parts = []
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        raise ValueError(f"Erro ao ler PDF: {str(e)}")

    return '\n'.join(text_parts).strip()


def extract_text_from_docx(filepath: str) -> str:
    """
    Extrai texto de um arquivo DOCX usando python-docx.
    Inclui parágrafos e texto de tabelas.
    """
    if not HAS_DOCX:
        raise ImportError("python-docx não está instalado. Execute: pip install python-docx")

    text_parts = []
    try:
        doc = DocxDocument(filepath)

        # Extrai parágrafos normais
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extrai texto de tabelas
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

    except Exception as e:
        raise ValueError(f"Erro ao ler DOCX: {str(e)}")

    return '\n'.join(text_parts).strip()


def process_upload(file: FileStorage) -> dict:
    """
    Salva o arquivo enviado e extrai seu texto.

    Args:
        file: Objeto FileStorage do Werkzeug (request.files['arquivo'])

    Returns:
        {
          'success': True,
          'filename': 'nome_salvo.pdf',
          'filepath': '/caminho/absoluto/arquivo.pdf',
          'texto': 'Texto extraído...',
          'num_chars': 1234,
        }
        ou
        {
          'success': False,
          'error': 'Mensagem de erro'
        }
    """
    if not file or file.filename == '':
        return {'success': False, 'error': 'Nenhum arquivo selecionado.'}

    if not allowed_file(file.filename):
        return {
            'success': False,
            'error': 'Formato não suportado. Use PDF ou DOCX.'
        }

    # Gera nome único para evitar conflitos
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_name = f"{uuid.uuid4().hex}.{ext}"

    upload_folder = current_app.config['UPLOAD_FOLDER']
    os.makedirs(upload_folder, exist_ok=True)
    filepath = os.path.join(upload_folder, unique_name)

    try:
        file.save(filepath)
    except Exception as e:
        return {'success': False, 'error': f'Erro ao salvar arquivo: {str(e)}'}

    # Extrai texto conforme o tipo
    try:
        if ext == 'pdf':
            texto = extract_text_from_pdf(filepath)
        elif ext == 'docx':
            texto = extract_text_from_docx(filepath)
        else:
            texto = ''
    except (ImportError, ValueError) as e:
        # Remove o arquivo se não conseguir processar
        if os.path.exists(filepath):
            os.remove(filepath)
        return {'success': False, 'error': str(e)}

    if not texto:
        return {
            'success': False,
            'error': 'Não foi possível extrair texto do arquivo. Verifique se o PDF não é escaneado.'
        }

    return {
        'success': True,
        'filename': unique_name,
        'filepath': filepath,
        'texto': texto,
        'num_chars': len(texto),
    }
